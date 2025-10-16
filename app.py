import streamlit as st
import os
import tempfile
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re
from docx import Document
import fitz

# Helper functions for reading different file types
def read_docx(uploaded_file):
    """
    Reads content from a Streamlit UploadedFile object for .docx files.

    Args:
        uploaded_file (streamlit.uploaded_file.UploadedFile): The uploaded .docx file object.

    Returns:
        str: The text content of the .docx file.
    """
    # Use the UploadedFile object directly
    doc = Document(uploaded_file)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def read_pdf(uploaded_file):
    """
    Reads content from a Streamlit UploadedFile object for .pdf files.

    Args:
        uploaded_file (streamlit.uploaded_file.UploadedFile): The uploaded .pdf file object.

    Returns:
        str: The text content of the .pdf file.
    """
    # Create a temporary file to read the content for PyMuPDF
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    try:
        doc = fitz.open(tmp_path)
        text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text()
        return text
    finally:
        os.remove(tmp_path) # Clean up the temporary file


# Data loading and preprocessing function
def load_and_preprocess_transcripts(uploaded_files):
    """
    Loads and preprocesses meeting transcripts from uploaded Streamlit files.

    Args:
        uploaded_files (list): A list of Streamlit UploadedFile objects.

    Returns:
        list: A list of preprocessed transcript strings.
    """
    transcripts = []
    for uploaded_file in uploaded_files:
        try:
            if uploaded_file.name.endswith(".txt"):
                 # Decode directly from the BytesIO object
                transcript_content = uploaded_file.getvalue().decode('utf-8')
                transcripts.append(transcript_content)
            elif uploaded_file.name.endswith(".docx"):
                transcript_content = read_docx(uploaded_file)
                transcripts.append(transcript_content)
            elif uploaded_file.name.endswith(".pdf"):
                transcript_content = read_pdf(uploaded_file)
                transcripts.append(transcript_content)
            else:
                st.warning(f"Skipping unsupported file type: {uploaded_file.name}")
        except Exception as e:
            st.error(f"Error processing file {uploaded_file.name}: {e}")


    cleaned_transcripts = []
    for transcript_content in transcripts:
        # Clean and normalize the text
        cleaned_content = re.sub(r'[^\w\s.]', '', transcript_content) # Remove characters except alphanumeric, whitespace, and periods
        cleaned_content = cleaned_content.lower() # Convert to lowercase
        cleaned_transcripts.append(cleaned_content)

    return cleaned_transcripts

# Chunking and embedding function
def chunk_and_embed(preprocessed_transcripts):
    """
    Splits and embeds preprocessed transcript strings.

    Args:
        preprocessed_transcripts (list): A list of preprocessed transcript strings.

    Returns:
        tuple: A tuple containing:
            - list: A list of document chunks.
            - list: A list of corresponding embeddings.
    """
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    document_chunks = text_splitter.create_documents(preprocessed_transcripts)

    embedding_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    document_embeddings = embedding_model.embed_documents([chunk.page_content for chunk in document_chunks])

    return document_chunks, document_embeddings

# Vector store creation function
def create_vector_store(document_chunks, document_embeddings):
    """
    Creates a Chroma vector store from document chunks and embeddings.

    Args:
        document_chunks (list): A list of document chunks.
        document_embeddings (list): A list of corresponding embeddings.

    Returns:
        Chroma: The created vector store object.
    """
    embeddings_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    # Using an in-memory vector store for the Streamlit demo
    vectorstore = Chroma.from_documents(document_chunks, embeddings_model)
    return vectorstore

# Streamlit application layout
st.title("Meeting Transcript RAG Pipeline")

st.header("Upload Meeting Transcripts")
uploaded_files = st.file_uploader("Choose transcript files (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    # Process uploaded files only if different from previous uploads
    if "uploaded_file_names" not in st.session_state or st.session_state.uploaded_file_names != [f.name for f in uploaded_files]:
        st.session_state.uploaded_file_names = [f.name for f in uploaded_files]
        with st.spinner("Processing transcripts..."):
            preprocessed_data = load_and_preprocess_transcripts(uploaded_files)
            if preprocessed_data: # Only proceed if data was successfully processed
                document_chunks, document_embeddings = chunk_and_embed(preprocessed_data)
                st.session_state.vectorstore = create_vector_store(document_chunks, document_embeddings)
                st.success("Transcripts processed and vector store created!")
            else:
                st.warning("No valid transcripts were processed.")


    st.header("Ask a Question")
    user_question = st.text_input("Enter your question:")

    if user_question and "vectorstore" in st.session_state:
        with st.spinner("Getting answer..."):
            # Build RAG chain (re-instantiate LLM and retriever as needed by Streamlit's execution model)
            # Assuming OPENAI_API_KEY is set as an environment variable
            # os.environ["OPENAI_API_KEY"] = userdata.get("OPEN_API_KEY") # Modified to use os.environ
            # Attempt to read the API key from environment variables
            openai_api_key = os.environ.get("OPENAI_API_KEY")
            if not openai_api_key:
                st.error("OPENAI_API_KEY environment variable not set. Please set it in your deployment environment.")
            else:
                llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, api_key=openai_api_key)
                retriever = st.session_state.vectorstore.as_retriever()
                rag_chain = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)

                answer = rag_chain.invoke({"query": user_question})
                st.write("Answer:")
                st.info(answer['result'])
    elif user_question and "vectorstore" not in st.session_state:
        st.warning("Please upload and process transcripts first.")
else:
    st.info("Please upload meeting transcripts to get started.")






# import streamlit as st
# import os
# import tempfile
# from langchain.chains import RetrievalQA
# from langchain_openai import ChatOpenAI
# from langchain_community.vectorstores import Chroma
# from langchain_community.embeddings import SentenceTransformerEmbeddings
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import re
# from docx import Document
# import fitz
# # from google.colab import userdata

# # Helper functions for reading different file types
# def read_docx(filepath):
#     st.write("Reading Documents")
#     doc = Document(filepath)
#     text = []
#     for paragraph in doc.paragraphs:
#         text.append(paragraph.text)
#     return '\n'.join(text)

# def read_pdf(filepath):
#     st.write("Reading PDFs")
#     doc = fitz.open(filepath)
#     text = ""
#     for page_num in range(doc.page_count):
#         page = doc.load_page(page_num)
#         text += page.get_text()
#     return text

# # Data loading and preprocessing function
# def load_and_preprocess_transcripts(uploaded_files):
#     """
#     Loads and preprocesses meeting transcripts from uploaded Streamlit files.

#     Args:
#         uploaded_files (list): A list of Streamlit UploadedFile objects.

#     Returns:
#         list: A list of preprocessed transcript strings.
#     """
#     st.write("Loading and Pre-processing Transcripts")
#     transcripts = []
#     for uploaded_file in uploaded_files:
#         # Create a temporary file to read the content
#         with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
#             tmp_file.write(uploaded_file.getvalue())
#             tmp_path = tmp_file.name

    #     try:
    #         if uploaded_file.name.endswith(".txt"):
    #             with open(tmp_path, 'r', encoding='utf-8') as f:
    #                 transcript_content = f.read()
    #                 transcripts.append(transcript_content)
    #         elif uploaded_file.name.endswith(".docx"):
    #             transcript_content = read_docx(tmp_path)
    #             transcripts.append(transcript_content)
    #         elif uploaded_file.name.endswith(".pdf"):
    #             transcript_content = read_pdf(tmp_path)
    #             transcripts.append(transcript_content)
    #         else:
    #             st.warning(f"Skipping unsupported file type: {uploaded_file.name}")
    #     finally:
    #         os.remove(tmp_path) # Clean up the temporary file


    # cleaned_transcripts = []
    # for transcript_content in transcripts:
    #     # Clean and normalize the text
    #     cleaned_content = re.sub(r'[^\w\s.]', '', transcript_content) # Remove characters except alphanumeric, whitespace, and periods
    #     cleaned_content = cleaned_content.lower() # Convert to lowercase
    #     cleaned_transcripts.append(cleaned_content)

    # return cleaned_transcripts

# Chunking and embedding function
def chunk_and_embed(preprocessed_transcripts):
    """
    Splits and embeds preprocessed transcript strings.

    Args:
        preprocessed_transcripts (list): A list of preprocessed transcript strings.

    Returns:
        tuple: A tuple containing:
            - list: A list of document chunks.
            - list: A list of corresponding embeddings.
    """
    st.write("Splitting, Chunking and Embedding Transcript Strings")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    document_chunks = text_splitter.create_documents(preprocessed_transcripts)

    embedding_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    document_embeddings = embedding_model.embed_documents([chunk.page_content for chunk in document_chunks])

    return document_chunks, document_embeddings

# Vector store creation function
def create_vector_store(document_chunks, document_embeddings):
    """
    Creates a Chroma vector store from document chunks and embeddings.

    Args:
        document_chunks (list): A list of document chunks.
        document_embeddings (list): A list of corresponding embeddings.

    Returns:
        Chroma: The created vector store object.
    """
    st.write("Creating Vector Store")
    embeddings_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    # Using an in-memory vector store for the Streamlit demo
    vectorstore = Chroma.from_documents(document_chunks, embeddings_model)
    return vectorstore

# Streamlit application layout
st.title("BB RAG Pipeline for Meeting Transcripts")

st.header("Upload Meeting Transcripts")
uploaded_files = st.file_uploader("Choose transcript files (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    # Process uploaded files only if different from previous uploads
    if "uploaded_file_names" not in st.session_state or st.session_state.uploaded_file_names != [f.name for f in uploaded_files]:
        st.session_state.uploaded_file_names = [f.name for f in uploaded_files]
        with st.spinner("Processing transcripts..."):
            preprocessed_data = load_and_preprocess_transcripts(uploaded_files)
            document_chunks, document_embeddings = chunk_and_embed(preprocessed_data)
            st.session_state.vectorstore = create_vector_store(document_chunks, document_embeddings)
        st.success("Transcripts processed and vector store created!")

    st.header("Ask a Question")
    user_question = st.text_input("Enter your question:")

    if user_question and "vectorstore" in st.session_state:
        with st.spinner("Getting answer..."):
            # Build RAG chain (re-instantiate LLM and retriever as needed by Streamlit's execution model)
            # Assuming OPENAI_API_KEY is set as an environment variable
            # os.environ["OPENAI_API_KEY"] = userdata.get("OPEN_API_KEY")

            llm = ChatOpenAI(openai_api_key=os.environ["OPENAI_API_KEY"],model="gpt-4o-mini", temperature=0)
            retriever = st.session_state.vectorstore.as_retriever()
            rag_chain = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)

            answer = rag_chain.invoke({"query": user_question})
            st.write("Answer:")
            st.info(answer['result'])
    elif user_question and "vectorstore" not in st.session_state:
        st.warning("Please upload and process transcripts first.")
else:
    st.info("Please upload meeting transcripts to get started.")
